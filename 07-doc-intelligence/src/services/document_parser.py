"""PDF, DOCX, and TXT document parser.

Extracts raw text and metadata from uploaded files using PyMuPDF for
PDFs, python-docx for DOCX files, and plain-text decoding for TXT.
"""

import re
from typing import Any

import fitz  # PyMuPDF


class DocumentParser:
    """Parse uploaded documents into plain text and metadata.

    Supports PDF (via PyMuPDF), DOCX (via python-docx), and plain
    text files.  The :meth:`parse` entry-point routes to the correct
    parser based on file extension.
    """

    def parse(self, file_bytes: bytes, filename: str) -> tuple[str, dict[str, Any]]:
        """Parse *file_bytes* and return ``(full_text, metadata_dict)``.

        Args:
            file_bytes: Raw file content.
            filename: Original filename (used for extension routing).

        Raises:
            ValueError: If the file extension is not supported.

        Returns:
            A tuple of extracted text and a metadata dictionary.
        """
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return self._parse_pdf(file_bytes)
        if lower.endswith(".txt"):
            return self._parse_txt(file_bytes)
        if lower.endswith(".docx"):
            return self._parse_docx(file_bytes)
        raise ValueError(f"Unsupported file type: {filename}")

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_bytes: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from a PDF using PyMuPDF.

        Args:
            file_bytes: Raw PDF bytes.

        Returns:
            ``(text, metadata)`` where metadata contains page count
            and detected sections.
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = doc.page_count
        all_text_parts: list[str] = []

        for page in doc:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            spans: list[dict[str, Any]] = []
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        spans.append({
                            "text": span.get("text", ""),
                            "bbox": span.get("bbox", (0, 0, 0, 0)),
                            "font_size": span.get("size", 0),
                            "flags": span.get("flags", 0),
                        })

            spans.sort(key=lambda s: (round(s["bbox"][1] / 10), s["bbox"][0]))

            page_text = " ".join(s["text"] for s in spans if s["text"].strip())
            all_text_parts.append(page_text)

        full_text = "\n\n".join(all_text_parts)
        sections = self._detect_sections(full_text)
        language = self._detect_language(full_text)

        metadata: dict[str, Any] = {
            "page_count": page_count,
            "sections": sections,
            "language": language,
        }
        doc.close()
        return full_text, metadata

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _parse_txt(self, file_bytes: bytes) -> tuple[str, dict[str, Any]]:
        """Decode raw bytes as a plain text file.

        Tries UTF-8 first, then falls back to latin-1 which never fails.

        Args:
            file_bytes: Raw text file bytes.

        Returns:
            ``(text, metadata)``.
        """
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text = file_bytes.decode("latin-1")

        sections = self._detect_sections(text)
        language = self._detect_language(text)

        metadata: dict[str, Any] = {
            "page_count": 1,
            "sections": sections,
            "language": language,
        }
        return text, metadata

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _parse_docx(self, file_bytes: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from a DOCX file using python-docx.

        Args:
            file_bytes: Raw DOCX bytes.

        Returns:
            ``(text, metadata)``.
        """
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs)
        sections = self._detect_sections(full_text)
        language = self._detect_language(full_text)

        metadata: dict[str, Any] = {
            "page_count": 1,
            "sections": sections,
            "language": language,
        }
        return full_text, metadata

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _detect_sections(self, text: str) -> list[dict[str, str]]:
        """Detect section headers via numbered-heading patterns.

        Args:
            text: Full document text.

        Returns:
            List of dicts with ``title`` and ``level`` keys.
        """
        patterns = [
            (r"^(\d+\.\d+\.\d+)\s+(.+)$", "3"),
            (r"^(\d+\.\d+)\s+(.+)$", "2"),
            (r"^(\d+)\.\s+(.+)$", "1"),
            (r"^[Ss]ection\s+(\d+)\s*:?\s*(.+)$", "1"),
            (r"^[A-Z][A-Z\s]{5,}$", "1"),
        ]
        sections: list[dict[str, str]] = []
        seen: set[str] = set()

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            for regex, level in patterns:
                match = re.match(regex, stripped)
                if match:
                    title = stripped
                    if title not in seen:
                        seen.add(title)
                        sections.append({"title": title, "level": level})
                    break

        return sections

    def _detect_language(self, text: str) -> str:
        """Detect document language by simple word-frequency heuristics.

        Args:
            text: Document text to analyse.

        Returns:
            ISO 639-1 language code (defaults to ``"en"``).
        """
        sample = text[:5000].lower()
        english_words = {"the", "and", "is", "of", "to", "in", "for", "that", "with", "this"}
        spanish_words = {"el", "la", "de", "en", "que", "y", "los", "del", "las", "por"}
        german_words = {"der", "die", "und", "ist", "von", "den", "das", "nicht", "ein", "auf"}

        words = set(re.findall(r"\b[a-zäöüß]{3,}\b", sample))

        en_score = len(words & english_words)
        es_score = len(words & spanish_words)
        de_score = len(words & german_words)

        best = max(en_score, es_score, de_score)
        if best == 0:
            return "en"
        if en_score == best:
            return "en"
        if es_score == best:
            return "es"
        return "de"
