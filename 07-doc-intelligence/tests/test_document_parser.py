"""Tests for the document parser module."""

import pytest

from src.services.document_parser import DocumentParser


@pytest.fixture
def parser() -> DocumentParser:
    """Return a DocumentParser instance."""
    return DocumentParser()


class TestParseTxt:
    """Tests for plain-text parsing."""

    def test_parse_txt_utf8(self, parser: DocumentParser) -> None:
        """UTF-8 text is decoded correctly."""
        content = "Section 1: Safety Data Sheet\n\nThis document contains safety information."
        text, meta = parser._parse_txt(content.encode("utf-8"))
        assert "Safety Data Sheet" in text
        assert meta["page_count"] == 1

    def test_parse_txt_latin1(self, parser: DocumentParser) -> None:
        """Latin-1 encoded text is handled gracefully."""
        content = "Café résumé naïve"
        text, meta = parser._parse_txt(content.encode("latin-1"))
        assert "Café" in text

    def test_parse_txt_empty(self, parser: DocumentParser) -> None:
        """Empty file returns empty text."""
        text, meta = parser._parse_txt(b"")
        assert text == ""
        assert meta["page_count"] == 1

    def test_parse_txt_routes_correctly(self, parser: DocumentParser) -> None:
        """The parse() method routes .txt files to _parse_txt."""
        text, meta = parser.parse(b"Hello world", "test.txt")
        assert text == "Hello world"


class TestDetectSections:
    """Tests for section header detection."""

    def test_numbered_sections(self, parser: DocumentParser) -> None:
        """Detects numbered section headers."""
        text = "1. Introduction\nSome text\n2. Hazard Identification\nMore text"
        sections = parser._detect_sections(text)
        titles = [s["title"] for s in sections]
        assert "1. Introduction" in titles
        assert "2. Hazard Identification" in titles

    def test_subsections(self, parser: DocumentParser) -> None:
        """Detects subsection headers (x.y format)."""
        text = "1.1 Physical Properties\nDetails here"
        sections = parser._detect_sections(text)
        assert len(sections) == 1
        assert sections[0]["level"] == "2"

    def test_no_sections(self, parser: DocumentParser) -> None:
        """Returns empty list when no sections are found."""
        text = "Just a paragraph with no headers at all."
        sections = parser._detect_sections(text)
        assert sections == []

    def test_section_keyword(self, parser: DocumentParser) -> None:
        """Detects 'Section X' headers."""
        text = "Section 3: First Aid Measures\nContent here"
        sections = parser._detect_sections(text)
        assert any("Section 3" in s["title"] for s in sections)


class TestDetectLanguage:
    """Tests for language detection heuristic."""

    def test_english(self, parser: DocumentParser) -> None:
        """Identifies English text."""
        text = "The chemical is the substance of this and that for the work"
        lang = parser._detect_language(text)
        assert lang == "en"

    def test_spanish(self, parser: DocumentParser) -> None:
        """Identifies Spanish text."""
        text = "el la de en que y los del las por el la de en que"
        lang = parser._detect_language(text)
        assert lang == "es"

    def test_empty(self, parser: DocumentParser) -> None:
        """Defaults to English for empty text."""
        lang = parser._detect_language("")
        assert lang == "en"


class TestParsePdf:
    """Tests for PDF parsing using synthetic bytes."""

    def test_parse_pdf_basic(self, parser: DocumentParser) -> None:
        """Parses a minimal valid PDF."""
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hazard Identification\nSodium Hydroxide 50%")
        pdf_bytes = doc.tobytes()
        doc.close()
        text, meta = parser._parse_pdf(pdf_bytes)
        assert "Sodium Hydroxide" in text
        assert meta["page_count"] == 1

    def test_parse_pdf_multi_page(self, parser: DocumentParser) -> None:
        """Handles multi-page PDFs."""
        import fitz
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "Page one content")
        doc.new_page().insert_text((72, 72), "Page two content")
        pdf_bytes = doc.tobytes()
        doc.close()
        text, meta = parser._parse_pdf(pdf_bytes)
        assert meta["page_count"] == 2
        assert "Page one" in text
        assert "Page two" in text


class TestParseDocx:
    """Tests for DOCX parsing using synthetic bytes."""

    def test_parse_docx_basic(self, parser: DocumentParser) -> None:
        """Parses a DOCX with paragraphs and a table."""
        import io
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("First paragraph about chemicals")
        doc.add_paragraph("Second paragraph about safety")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Chemical"
        table.cell(0, 1).text = "Concentration"
        table.cell(1, 0).text = "Methanol"
        table.cell(1, 1).text = "50%"
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        text, meta = parser._parse_docx(docx_bytes)
        assert "chemicals" in text.lower()
        assert "Methanol" in text
        assert meta["page_count"] == 1


class TestUnsupportedExtension:
    """Tests for unsupported file types."""

    def test_raises_on_unsupported(self, parser: DocumentParser) -> None:
        """Raises ValueError for unknown extensions."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            parser.parse(b"data", "file.xyz")
